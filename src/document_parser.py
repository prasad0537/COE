from __future__ import annotations

import base64
import binascii
import io
import re
from pathlib import Path


MAX_DOCUMENT_SIZE_BYTES = 10 * 1024 * 1024
SUPPORTED_DOCUMENT_TYPES = {
    ".txt": "Plain Text",
    ".pdf": "PDF",
    ".docx": "Word Document",
}


def supported_upload_extensions() -> str:
    return ", ".join(extension.lstrip(".").upper() for extension in SUPPORTED_DOCUMENT_TYPES)


def normalize_text(text: str) -> str:
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def decode_payload(content_base64: str) -> bytes:
    try:
        return base64.b64decode(content_base64, validate=True)
    except (ValueError, binascii.Error) as exc:
        raise ValueError("The uploaded file payload is not valid base64 data.") from exc


def extract_text_from_txt(file_bytes: bytes) -> str:
    for encoding in ("utf-8-sig", "utf-16", "utf-16-le", "utf-16-be", "latin-1"):
        try:
            return file_bytes.decode(encoding)
        except UnicodeDecodeError:
            continue
    return file_bytes.decode("utf-8", errors="ignore")


def extract_text_from_pdf(file_bytes: bytes) -> str:
    try:
        from pypdf import PdfReader
    except ImportError as exc:  # pragma: no cover
        raise RuntimeError("PDF support is not installed on the server.") from exc

    reader = PdfReader(io.BytesIO(file_bytes))
    return "\n\n".join((page.extract_text() or "").strip() for page in reader.pages)


def extract_text_from_docx(file_bytes: bytes) -> str:
    try:
        from docx import Document
    except ImportError as exc:  # pragma: no cover
        raise RuntimeError("DOCX support is not installed on the server.") from exc

    document = Document(io.BytesIO(file_bytes))
    blocks: list[str] = []

    for paragraph in document.paragraphs:
        if paragraph.text.strip():
            blocks.append(paragraph.text.strip())

    for table in document.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                blocks.append(row_text)

    return "\n\n".join(blocks)


def extract_document_text(filename: str, content_base64: str) -> dict[str, object]:
    safe_filename = Path(filename or "").name
    extension = Path(safe_filename).suffix.lower()

    if not safe_filename:
        raise ValueError("The uploaded file is missing a filename.")

    if extension == ".doc":
        raise ValueError("Legacy .doc files are not supported yet. Convert the file to .docx or PDF and try again.")

    if extension not in SUPPORTED_DOCUMENT_TYPES:
        supported = supported_upload_extensions()
        raise ValueError(f"Unsupported file type. Upload one of: {supported}.")

    file_bytes = decode_payload(content_base64)
    if not file_bytes:
        raise ValueError("The uploaded file is empty.")

    if len(file_bytes) > MAX_DOCUMENT_SIZE_BYTES:
        raise ValueError("The uploaded file is too large. Use a file up to 10 MB.")

    if extension == ".txt":
        extracted = extract_text_from_txt(file_bytes)
    elif extension == ".pdf":
        extracted = extract_text_from_pdf(file_bytes)
    else:
        extracted = extract_text_from_docx(file_bytes)

    normalized = normalize_text(extracted)
    if not normalized:
        raise ValueError("No readable text was found in the uploaded file.")

    return {
        "filename": safe_filename,
        "file_type": SUPPORTED_DOCUMENT_TYPES[extension],
        "text": normalized,
        "character_count": len(normalized),
    }
